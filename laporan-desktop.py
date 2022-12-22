from docx import Document
from docx.shared import Inches, Pt, Length, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from datetime import date
import datetime

#variabel
doc = Document()
style = doc.styles
data_tabel_bawah = []
n = [0]
m = len(n)
name = str(input("Masukan Nama Lengkap Anda Beserta Gelar: "))
jabatan = str(input("Masukan Jabatan Anda: "))

#------Fungsi save dokumen
def save():
    doc.save(name+".docx")

#------Fungsi perulangan dan compile dokumen
def perulangan():
    userInput = input("Apakah ingin menambah data? (y/n) ")
    if userInput == 'y' or userInput == 'Y': 
        for i in range(1): 
            n.append((str(i+1)))
        tambahData()
    elif userInput == 'n' or userInput == 'N': 
        for no, target, tanggal, tercapai, alasan, solusi in data_tabel_bawah:
            b_tBawah = tabel_bawah.add_row().cells
            b_tBawah[0].text = no
            b_tBawah[1].text = target
            b_tBawah[2].text = tanggal
            b_tBawah[3].text = tercapai
            b_tBawah[4].text = alasan
            b_tBawah[5].text = solusi
        styleTeks("")
        styleTeks("Masukan/Usulan Program Inovasi/Kreativitas: ")
        
        #tabel usulan
        tabel_usulan = doc.add_table(rows=1, cols=1)
        tabel_usulan.style = 'Table Grid'
        isi_tabel_usulan = tabel_usulan.rows[0].cells
        isi_tabel_usulan[0].text = input("Masukan/Usulan Program Inovasi/Kreativitas: ")
        isi_tabel_usulan[0].height = Inches(2)

        styleTeks("")
        waktuNih()
        styleTeks("Dibuat oleh: ")
        styleTeks("")
        styleTeks("")
        styleTeks2(f"({name})")
        styleTeks2(jabatan)
        save()
        input(f"Terima kasih {name} sudah menggunakan produk Ariftiarno System")
    elif userInput != 'n' or userInput != 'N' or userInput != 'y' or userInput != 'Y':
        print("Tombol yang Anda ketik tidak terbaca oleh Sistem. Silakan ketik huruf Y atau N untuk melanjutkan")
        perulangan()

#------Fungsi tambah data ke tabel bawah
def tambahData():
    for i in range(m):
        t_no = str(len(data_tabel_bawah)+1)
        t_trg = input("Masukan Target Kegiatan: ")
        t_tgl = input("Masukan Tanggal: ")
        t_tcp = input("Tercapai atau Tidak? ")
        t_als = input("Masukan Alasan: ")
        t_sol = input("Masukan Solusi: ")
        data_tabel_bawah.append([t_no, t_trg, t_tgl, t_tcp, t_als, t_sol])
        perulangan()

#------Fungsi Style
def change_orientation():
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_height = Mm(210)
    section.page_width = Mm(297)
    section.left_margin = Mm(25.4)
    section.right_margin = Mm(25.4)
    section.top_margin = Mm(25.4)
    section.bottom_margin = Mm(25.4)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)

def styleTeks(a):
    teksPara = doc.add_paragraph().add_run(a)
    teksPara.font.size = Pt(12)
    teksPara.font.name = 'Calibri'
    return teksPara

def styleTeks2(a):
    teksPara = doc.add_paragraph()
    teksPara.paragraph_format.space_after = Pt(0)
    teksPara_font = teksPara.add_run(a)
    teksPara_font.font.size = Pt(12)
    teksPara_font.font.name = 'Calibri'
    return teksPara

def styleJudul(a):
    teksJudul = doc.add_paragraph()
    teksJudul.alignment = WD_ALIGN_PARAGRAPH.CENTER
    teksJudul.paragraph_format.space_after = Pt(0)
    teksJudul_font = teksJudul.add_run(a)
    teksJudul_font.font.name = 'Calibri'
    teksJudul_font.font.size = Pt(18)
    return teksJudul

#------Fungsi waktu dan tanggal
def waktuNih():
    waktu = datetime.datetime.now()
    bulan = waktu.strftime("%m")
    hari = waktu.strftime("%d")
    tahun = waktu.strftime("%Y")
    if bulan == "1":
        styleTeks(f"Bogor, {hari} Januari {tahun}")
    elif bulan == "2":
        styleTeks(f"Bogor, {hari} Februari {tahun}")
    elif bulan == "3":
        styleTeks(f"Bogor, {hari} Maret {tahun}")
    elif bulan == "4":
        styleTeks(f"Bogor, {hari} April {tahun}")
    elif bulan == "5":
        styleTeks(f"Bogor, {hari} Mei {tahun}")
    elif bulan == "6":
        styleTeks(f"Bogor, {hari} Juni {tahun}")
    elif bulan == "7":
        styleTeks(f"Bogor, {hari} Juli {tahun}")
    elif bulan == "8":
        styleTeks(f"Bogor, {hari} Agustus {tahun}")
    elif bulan == "9":
        styleTeks(f"Bogor, {hari} September {tahun}")
    elif bulan == "10":
        styleTeks(f"Bogor, {hari} Oktober {tahun}")
    elif bulan == "11":
        styleTeks(f"Bogor, {hari} November {tahun}")
    elif bulan == "12":
        styleTeks(f"Bogor, {hari} Desember {tahun}")
    else:
        print("Gagal Mencetak Tanggal")

#------Program dimulai dari sini------#
change_orientation()
styleJudul("LAPORAN BULANAN")
styleJudul("UNIT FAKULTAS EKONOMI DAN BISNIS")
styleJudul("UNIVERSITAS PAKUAN")
styleJudul("")

#------Tabel atas
tabel_atas = doc.add_table(rows=3, cols=2)
tabel_atas.style = 'Table Grid'
judulTab_atas = tabel_atas.columns[0].cells
for cell in judulTab_atas:
    cell.width = Inches(0.73)
    cell.height = Inches(0.21)
judulTab_atas[0].text = 'Unit'
judulTab_atas[1].text = 'Bulan'
judulTab_atas[2].text = 'Tahun'
isiTab_atas = tabel_atas.columns[1].cells
for cell in isiTab_atas:
    cell.width = Inches(3)
    cell.height = Inches(0.21)
isiTab_atas[0].text = input("Dari Unit Mana? ")
isiTab_atas[1].text = input("Laporan Bulan Apa? ")
isiTab_atas[2].text = input("Tahun? ")

styleJudul("")

#------Tabel bawah
tabel_bawah = doc.add_table(rows=1, cols=6)
tabel_bawah.style = 'Table Grid'
judulTab_bawah = tabel_bawah.rows[0].cells
for cell in judulTab_bawah:
    cell.height = Inches(0.21)
judulTab_bawah[0].text = 'No'
judulTab_bawah[0].width = Inches(0)
judulTab_bawah[1].text = 'Target Kegiatan'
judulTab_bawah[1].width = Inches(3)
judulTab_bawah[2].text = 'Tanggal'
judulTab_bawah[3].text = 'Tercapai/Tidak'
judulTab_bawah[4].text = 'Alasan'
judulTab_bawah[5].text = 'Solusi'
tambahData()






